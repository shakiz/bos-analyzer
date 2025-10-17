from fastapi import FastAPI, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import pandas as pd
import re
from io import BytesIO
from typing import List, Tuple

app = FastAPI()

# Allow frontend access
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -------------------
# Helpers
# -------------------
def clean_text(s: str) -> str:
    s = s.replace("\xa0", " ").replace("\u200b", "").strip()
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\r\n?", "\n", s)
    return s

def normalize_phone_candidates(text: str) -> List[str]:
    candidates = re.findall(r"(?:\+?\d[\d\s\-\(\)]{7,}\d)", text)
    phones = []
    for c in candidates:
        digits = re.sub(r"\D", "", c)
        if 9 < len(digits) <= 13:
            phones.append(digits)
    more = re.findall(r"\b\d{10,13}\b", text)
    for m in more:
        if m not in phones:
            phones.append(m)
    return phones

def eval_bill_expression(expr: str) -> float:
    if not expr or not isinstance(expr, str):
        raise ValueError("empty expr")
    expr_clean = expr.replace("taka", "").replace("tk", "").replace("Taka", "").replace(",", " ").strip()
    if "=" in expr_clean:
        rhs = expr_clean.split("=")[-1]
        nums = re.findall(r"\d+[.,]?\d*", rhs)
        if nums:
            return float(nums[-1].replace(",", ""))
    if re.search(r"[+\-*/]", expr_clean):
        safe = re.sub(r"[^\d+\-*/.() ]", "", expr_clean)
        if re.fullmatch(r"[0-9+\-*/.\s()]+", safe):
            try:
                val = eval(safe, {"__builtins__": None}, {})
                return float(val)
            except Exception:
                pass
    nums = re.findall(r"\d+[.,]?\d*", expr_clean)
    if nums:
        return float(nums[-1].replace(",", ""))
    raise ValueError("Could not parse amount from: " + expr)

def find_bill_lines(text: str) -> List[Tuple[int, str]]:
    lines = text.splitlines()
    results = []
    for i, line in enumerate(lines):
        if re.search(r"(?:total\s*bill|bill)[:\-]?", line, re.IGNORECASE):
            results.append((i, line.strip()))
    return results

# -------------------
# Main extraction
# -------------------
def extract_info_from_doc(doc_bytes, filename=None):
    document = Document(BytesIO(doc_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip() != ""]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip() != "":
                    paragraphs.append(cell.text)

    raw_text = "\n".join(paragraphs)
    raw_text = clean_text(raw_text)

    blocks = re.split(r"(?i)(?=^Name[:\-\s])", raw_text, flags=re.MULTILINE)
    if len(blocks) <= 1:
        blocks = [b.strip() for b in re.split(r"\n{2,}", raw_text) if b.strip()]

    size_counts = []
    customer_numbers = []
    customer_files = []
    order_amounts = []

    doc_phones = normalize_phone_candidates(raw_text)
    bill_lines = find_bill_lines(raw_text)

    for block in blocks:
        b = block.strip()
        if not b:
            continue

        # --- Sizes ---
        block_size_lines = [line for line in b.splitlines() if re.search(r"\bsize\b", line, re.IGNORECASE)]
        for line in block_size_lines:
            nums = re.findall(r"\d{2,3}", line)
            size_counts.extend(nums)

        # --- Customer numbers ---
        block_customer_lines = [line for line in b.splitlines() if re.search(r"(number|phone)", line, re.IGNORECASE)]
        for line in block_customer_lines:
            phones = normalize_phone_candidates(line)
            for ph in phones:
                customer_numbers.append(ph)
                customer_files.append((ph, filename))

        # --- Bill extraction ---
        block_lines = b.splitlines()
        for line in block_lines:
            if re.search(r"(?:total\s*bill|bill)[:\-]?", line, re.IGNORECASE):
                m = re.search(r"(?:total\s*bill|bill)[:\-]?\s*(.*)", line, re.IGNORECASE)
                expr = m.group(1).strip() if m else ""
                try:
                    amt = eval_bill_expression(expr)
                    phones_here = normalize_phone_candidates("\n".join(block_lines))
                    if phones_here:
                        for ph in phones_here:
                            order_amounts.append((ph, amt, filename))
                    elif doc_phones:
                        for ph in doc_phones:
                            order_amounts.append((ph, amt, filename))
                    else:
                        order_amounts.append(("unknown", amt, filename))
                except Exception:
                    pass

    # Fallback: attach any leftover bill lines to global phones
    if not order_amounts and bill_lines:
        for _, line in bill_lines:
            m = re.search(r"(?:total\s*bill|bill)[:\-]?\s*(.*)", line, re.IGNORECASE)
            if not m:
                continue
            expr = m.group(1)
            try:
                amt = eval_bill_expression(expr)
                if doc_phones:
                    for ph in doc_phones:
                        order_amounts.append((ph, amt, filename))
                else:
                    order_amounts.append(("unknown", amt, filename))
            except Exception:
                continue

    return size_counts, customer_numbers, order_amounts, customer_files


# -------------------
# API Endpoint
# -------------------
@app.post("/analyze")
async def analyze_files(files: List[UploadFile] = File(...)):
    all_sizes = []
    per_file = []
    all_customers = []
    all_customer_files = []
    all_order_amounts = []

    for file in files:
        content = await file.read()
        sizes, customers, order_amounts, customer_files = extract_info_from_doc(content, file.filename)
        all_sizes.extend(sizes)
        all_customers.extend(customers)
        all_order_amounts.extend(order_amounts)
        all_customer_files.extend(customer_files)

        df_file = pd.Series(sizes, name="size").value_counts().reset_index()
        df_file.columns = ["size", "count"]
        per_file.append({
            "filename": file.filename,
            "top_sizes": df_file.head(5).to_dict(orient="records"),
            "total_orders": len(sizes)
        })

    if not all_sizes:
        return {
            "top_sizes": [],
            "predicted_top_sizes": [],
            "total_orders": 0,
            "per_file": per_file,
            "top_customers": []
        }

    # --- Top sizes ---
    df = pd.Series(all_sizes, name="size").value_counts().reset_index()
    df.columns = ["size", "count"]
    top_sizes = df.head(5).to_dict(orient="records")

    # --- Predict Future Top Sizes (Golden Sizes) ---
    try:
        df_pred = df.copy()
        df_pred["size_num"] = pd.to_numeric(df_pred["size"], errors="coerce")
        df_pred = df_pred.dropna(subset=["size_num"])
        df_pred = df_pred.sort_values(by="size_num")
        df_pred["predicted_demand"] = (
            df_pred["count"].rolling(window=3, min_periods=1).mean() * 1.1
        )
        predicted_top_sizes = (
            df_pred.sort_values(by="predicted_demand", ascending=False)
            .head(5)
            .loc[:, ["size", "predicted_demand"]]
            .rename(columns={"predicted_demand": "count"})
            .to_dict(orient="records")
        )
    except Exception:
        predicted_top_sizes = []

    # --- Customers ---
    if all_customers:
        customer_counts = pd.Series(all_customers, name="customer").value_counts().reset_index()
        customer_counts.columns = ["customer", "order_count"]
        if len(files) == 1:
            customer_counts = customer_counts[customer_counts["order_count"] >= 1]
        else:
            customer_counts = customer_counts[customer_counts["order_count"] > 1]

        if all_order_amounts:
            df_amt = pd.DataFrame(all_order_amounts, columns=["customer", "amount", "filename"])
            amt_sum = df_amt.groupby("customer")["amount"].sum().reset_index()
            customer_counts = customer_counts.merge(amt_sum, on="customer", how="left")
            customer_counts["amount"] = customer_counts["amount"].fillna(0)
        else:
            customer_counts["amount"] = 0

        if all_customer_files:
            df_files = pd.DataFrame(all_customer_files, columns=["customer", "filename"])
            file_map = df_files.groupby("customer")["filename"].apply(
                lambda x: ", ".join(sorted(set(x)))
            ).reset_index()
            customer_counts = customer_counts.merge(file_map, on="customer", how="left")
        else:
            customer_counts["filename"] = ""

        top_customers = customer_counts.head(20).to_dict(orient="records")
    else:
        top_customers = []

    return {
        "top_sizes": top_sizes,
        "predicted_top_sizes": predicted_top_sizes,
        "total_orders": len(all_sizes),
        "per_file": per_file,
        "top_customers": top_customers,
    }
